import sys
import os
import subprocess
import platform
import unicodedata
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
from functools import partial
import multiprocessing as mp
from PyQt5.QtWidgets import (
    QApplication, QWidget, QPushButton, QLineEdit, QLabel,
    QVBoxLayout, QFileDialog, QListWidget, QRadioButton,
    QHBoxLayout, QProgressBar, QCheckBox, QMessageBox
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal, QTimer
import fitz  # PyMuPDF
import docx
import pytesseract
from pdf2image import convert_from_path
from tempfile import TemporaryDirectory
import gc

def find_tesseract():
    """Find Tesseract executable across platforms"""
    if platform.system() == "Windows":
        # Common Windows installation paths
        paths = [
            r"C:\Program Files\Tesseract-OCR\tesseract.exe",
            r"C:\tesseract\tesseract.exe",
            r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"
        ]
        for path in paths:
            if os.path.exists(path):
                return path
        # Try PATH
        try:
            subprocess.run(["tesseract", "--version"], capture_output=True, check=True)
            return "tesseract"
        except:
            pass
    else:
        # Linux/Unix - try common paths
        paths = ["/usr/bin/tesseract", "/usr/local/bin/tesseract"]
        for path in paths:
            if os.path.exists(path):
                return path
        # Try PATH
        try:
            subprocess.run(["tesseract", "--version"], capture_output=True, check=True)
            return "tesseract"
        except:
            pass
    return None

def find_poppler():
    """Find Poppler path across platforms"""
    if platform.system() == "Windows":
        # Common Windows installation paths
        paths = [
            r"C:\poppler\bin",
            r"C:\Program Files\poppler\bin",
            r"C:\Program Files (x86)\poppler\bin"
        ]
        for path in paths:
            if os.path.exists(os.path.join(path, "pdftoppm.exe")):
                return path
        # Try PATH
        try:
            subprocess.run(["pdftoppm", "-h"], capture_output=True, check=True)
            return None  # Available in PATH
        except:
            pass
    else:
        # Linux/Unix - usually in PATH
        try:
            subprocess.run(["pdftoppm", "-h"], capture_output=True, check=True)
            return None  # Available in PATH
        except:
            pass
    return None

# Set Tesseract path
TESSERACT_CMD = find_tesseract()
if TESSERACT_CMD:
    pytesseract.pytesseract.tesseract_cmd = TESSERACT_CMD

# Set Poppler path
POPPLER_PATH = find_poppler()

# Optimize for your system
MAX_WORKERS = min(32, (os.cpu_count() or 1) + 4)  # IO-bound threads
MAX_OCR_PROCESSES = max(1, (os.cpu_count() or 1) // 2)  # CPU-bound processes

def open_file(filepath):
    try:
        if platform.system() == "Windows":
            os.startfile(filepath)
        elif platform.system() == "Darwin":
            subprocess.run(["open", filepath], check=False)
        else:
            subprocess.run(["xdg-open", filepath], check=False)
    except Exception as e:
        print(f"Failed to open file: {e}")

def normalize_text(text):
    """Optimized text normalization with caching potential"""
    if not text:
        return ""
    text = unicodedata.normalize('NFKD', text.lower())
    return ''.join(c for c in text if not unicodedata.combining(c))

def process_single_file(file_args):
    """Standalone function for multiprocessing - processes a single file"""
    filepath, keyword, search_only_text = file_args
    try:
        if filepath.lower().endswith(".pdf"):
            return filepath, search_pdf_file(filepath, keyword, search_only_text)
        elif filepath.lower().endswith(".docx"):
            return filepath, search_docx_file(filepath, keyword)
        return filepath, False
    except Exception as e:
        print(f"[ERROR] Processing {filepath}: {e}")
        return filepath, False

def search_pdf_file(path, keyword, search_only_text=False):
    """Optimized PDF search with better memory management"""
    try:
        doc = fitz.open(path)
        has_text_pages = False
        norm_keyword = normalize_text(keyword)
        
        # First pass: check searchable text
        for page_num in range(len(doc)):
            page = doc[page_num]
            text = page.get_text()
            if text.strip():
                has_text_pages = True
                if norm_keyword in normalize_text(text):
                    doc.close()  # Clean up immediately
                    return True
            page = None  # Help GC
        
        doc.close()  # Clean up before OCR
        
        # Second pass: OCR if needed
        if not has_text_pages and not search_only_text:
            return search_pdf_with_ocr(path, norm_keyword)
            
        return False
        
    except Exception as e:
        print(f"[PDF ERROR] {path}: {e}")
        return False

def search_pdf_with_ocr(path, norm_keyword):
    """Separate OCR function for better memory management"""
    if not TESSERACT_CMD:
        print("[OCR ERROR] Tesseract not found")
        return False
        
    try:
        with TemporaryDirectory() as tmpdir:
            # Convert with lower DPI for faster processing
            images = convert_from_path(
                path, 
                dpi=150,  # Reduced from 200 for speed
                output_folder=tmpdir, 
                poppler_path=POPPLER_PATH,
                thread_count=2  # Limit poppler threads
            )
            
            for img in images:
                try:
                    # Use faster OCR settings
                    text = pytesseract.image_to_string(
                        img, 
                        lang='pol',
                        config='--psm 1 --oem 3'  # Optimized OCR settings
                    )
                    if norm_keyword in normalize_text(text):
                        return True
                except Exception as e:
                    print(f"[OCR ERROR]: {e}")
                finally:
                    img.close()  # Free image memory
                    img = None
            
        return False
        
    except Exception as e:
        print(f"[OCR PDF ERROR]: {e}")
        return False

def search_docx_file(path, keyword):
    """Optimized DOCX search"""
    try:
        doc = docx.Document(path)
        norm_keyword = normalize_text(keyword)
        
        for para in doc.paragraphs:
            if norm_keyword in normalize_text(para.text):
                return True
        
        # Also search in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if norm_keyword in normalize_text(cell.text):
                        return True
        
        return False
        
    except Exception as e:
        print(f"[DOCX ERROR] {path}: {e}")
        return False

class OptimizedFolderSearchWorker(QThread):
    progress = pyqtSignal(int)
    finished = pyqtSignal(list, int, int)
    file_processed = pyqtSignal(str, bool)  # For real-time updates

    def __init__(self, folder, keyword, search_only_text, max_workers=None):
        super().__init__()
        self.folder = folder
        self.keyword = keyword
        self.search_only_text = search_only_text
        self.results = []
        self.total = 0
        self.skipped = 0
        self.processed = 0
        self.max_workers = max_workers or MAX_WORKERS

    def run(self):
        # Collect all files first
        all_files = []
        for root, _, files in os.walk(self.folder):
            for name in files:
                if name.lower().endswith((".pdf", ".docx")):
                    all_files.append(os.path.join(root, name))

        self.total = len(all_files)
        if self.total == 0:
            self.finished.emit([], 0, 0)
            return

        # Prepare arguments for parallel processing
        file_args = [(filepath, self.keyword, self.search_only_text) 
                     for filepath in all_files]

        # Use ThreadPoolExecutor for I/O bound operations
        # For very large folders or OCR-heavy workloads, consider ProcessPoolExecutor
        with ThreadPoolExecutor(max_workers=self.max_workers) as executor:
            # Submit all tasks
            future_to_file = {
                executor.submit(process_single_file, args): args[0] 
                for args in file_args
            }
            
            # Process completed tasks
            for future in as_completed(future_to_file):
                filepath = future_to_file[future]
                try:
                    file_path, found = future.result()
                    if found:
                        self.results.append(file_path)
                        self.file_processed.emit(file_path, True)
                    
                    self.processed += 1
                    progress_percent = int((self.processed / self.total) * 100)
                    self.progress.emit(progress_percent)
                    
                except Exception as e:
                    print(f"[FUTURE ERROR] {filepath}: {e}")
                    self.skipped += 1
                    self.processed += 1

        self.finished.emit(self.results, self.total, self.skipped)

class KeywordSearcher(QWidget):
    def __init__(self):
        super().__init__()
        
        # Check dependencies at startup
        missing_deps = []
        if not TESSERACT_CMD:
            missing_deps.append("Tesseract OCR")
        if POPPLER_PATH is None and platform.system() == "Windows":
            # On Windows, we need explicit poppler path
            try:
                subprocess.run(["pdftoppm", "-h"], capture_output=True, check=True)
            except:
                missing_deps.append("Poppler")
        
        if missing_deps:
            self.show_dependency_warning(missing_deps)
        
        self.setWindowTitle("WYSZUKIWARKA CC 1.1")

        self.keyword_input = QLineEdit()
        self.keyword_input.setPlaceholderText("Wpisz słowo/słowa klucz...")
        self.keyword_input.setToolTip("Wpisz słowo lub frazę do wyszukania w dokumentach")

        self.folder_radio = QRadioButton("Szukaj w folderze")
        self.folder_radio.setToolTip("Przeszukaj wszystkie pliki PDF i DOCX w wybranym folderze")
        self.file_radio = QRadioButton("Szukaj w pojedynczym pliku")
        self.file_radio.setToolTip("Przeszukaj tylko jeden wybrany plik PDF lub DOCX")
        self.file_radio.setChecked(True)

        self.ocr_checkbox = QCheckBox("Szukaj tylko w wyszukiwalnych plikach PDF (pomiń OCR)")
        self.ocr_checkbox.setToolTip("Zaznacz aby pominąć skanowanie OCR - szybsze ale nie znajdzie tekstu w zeskanowanych PDF")

        # Add performance settings
        self.workers_label = QLabel(f"Równoczesne wątki: {MAX_WORKERS}")
        self.workers_label.setToolTip("Liczba równoczesnych wątków używanych do wyszukiwania")
        self.high_performance = QCheckBox("Tryb wysokiej wydajności (więcej RAM)")
        self.high_performance.setToolTip("Użyj więcej wątków dla szybszego wyszukiwania - wymaga więcej pamięci RAM")

        self.select_button = QPushButton("Wybierz plik/folder")
        self.select_button.setToolTip("Otwórz okno wyboru pliku lub folderu do przeszukania")
        self.search_button = QPushButton("Wyszukaj")
        self.search_button.setToolTip("Rozpocznij wyszukiwanie słowa kluczowego w wybranych plikach")
        self.stop_button = QPushButton("Zatrzymaj")
        self.stop_button.setToolTip("Zatrzymaj trwające wyszukiwanie")
        self.stop_button.setEnabled(False)
        self.about_button = QPushButton("O programie")
        self.about_button.setToolTip("Wyświetl informacje o programie")

        self.progress_bar = QProgressBar()
        self.progress_bar.setVisible(False)
        self.progress_bar.setToolTip("Postęp wyszukiwania w procentach")
        
        self.status_label = QLabel("")
        self.results_list = QListWidget()
        self.results_list.setToolTip("Lista znalezionych plików - kliknij dwukrotnie aby otworzyć plik")

        # Layout
        layout = QVBoxLayout()
        layout.addWidget(QLabel("Keyword:"))
        layout.addWidget(self.keyword_input)

        radio_layout = QHBoxLayout()
        radio_layout.addWidget(self.file_radio)
        radio_layout.addWidget(self.folder_radio)
        layout.addLayout(radio_layout)

        layout.addWidget(self.ocr_checkbox)
        layout.addWidget(self.high_performance)
        layout.addWidget(self.workers_label)
        
        button_layout = QHBoxLayout()
        button_layout.addWidget(self.select_button)
        button_layout.addWidget(self.search_button)
        button_layout.addWidget(self.stop_button)
        layout.addLayout(button_layout)
        
        layout.addWidget(self.progress_bar)
        layout.addWidget(self.status_label)
        layout.addWidget(QLabel("Wyniki:"))
        layout.addWidget(self.results_list)
        layout.addWidget(self.about_button)

        self.setLayout(layout)

        self.path = None
        self.worker = None

        # Connect signals
        self.select_button.clicked.connect(self.select_path)
        self.search_button.clicked.connect(self.search)
        self.stop_button.clicked.connect(self.stop_search)
        self.results_list.itemDoubleClicked.connect(self.open_item)
        self.about_button.clicked.connect(self.show_about)

    def show_dependency_warning(self, missing_deps):
        """Show warning about missing dependencies"""
        if platform.system() == "Windows":
            install_msg = (
                "Brakujące składniki:\n" + "\n".join([f"• {dep}" for dep in missing_deps]) +
                "\n\nInstrukcje instalacji Windows:\n"
                "• Tesseract: Pobierz z GitHub UB-Mannheim/tesseract\n"
                "  Zainstaluj w C:\\Program Files\\Tesseract-OCR\\\n"
                "• Poppler: Pobierz z blog.alivate.com.au/poppler-windows\n"
                "  Rozpakuj do C:\\poppler\\\n\n"
                "OCR (skanowane PDF) nie będzie działać bez tych składników."
            )
        else:
            install_msg = (
                "Brakujące składniki:\n" + "\n".join([f"• {dep}" for dep in missing_deps]) +
                "\n\nInstrukcje instalacji Linux:\n"
                "sudo apt update\n"
                "sudo apt install tesseract-ocr tesseract-ocr-pol poppler-utils\n\n"
                "OCR (skanowane PDF) nie będzie działać bez tych składników."
            )
        
        QMessageBox.warning(self, "Brakujące składniki", install_msg)

    def show_about(self):
        tesseract_status = "✓ Znaleziony" if TESSERACT_CMD else "✗ Nie znaleziony"
        poppler_status = "✓ Dostępny" if POPPLER_PATH is not None or platform.system() != "Windows" else "✗ Nie znaleziony"
        
        QMessageBox.information(
            self,
            "O programie",
            "WYSZUKIWARKA CC\n\n"
            "Aplikacja wyszukuje treść w zeskanowanych bądź wyszukiwalnych plikach PDF i DOCX.\n"
            "\n"
            f"Używa {MAX_WORKERS} wątków.\n"
            f"Status składników:\n"
            f"• Tesseract OCR: {tesseract_status}\n"
            f"• Poppler: {poppler_status}\n\n"
            "Zaprojektowane do wewnątrz firmowego użytku.\n"
            "Wszystkie prawa zastrzeżone\n"
            
            "Autor: Michał Skóra dla C&C Chakowski & Ciszek\n"
            "2025\n"
            "Wersja: 1.1"
        )

    def select_path(self):
        if self.folder_radio.isChecked():
            self.path = QFileDialog.getExistingDirectory(self, "Wybierz folder")
        else:
            self.path, _ = QFileDialog.getOpenFileName(self, "Wybierz plik", "", "Documents (*.pdf *.docx)")

    def search(self):
        self.results_list.clear()
        self.status_label.setText("")
        raw_keyword = self.keyword_input.text().strip()
        
        if not self.path or not raw_keyword:
            QMessageBox.warning(self, "Błąd", "Wybierz plik/folder i wpisz słowo kluczowe")
            return

        skip_ocr = self.ocr_checkbox.isChecked()
        self.search_button.setEnabled(False)
        self.stop_button.setEnabled(True)

        if self.folder_radio.isChecked():
            self.progress_bar.setVisible(True)
            self.progress_bar.setValue(0)
            
            # Adjust workers based on performance mode
            max_workers = MAX_WORKERS * 2 if self.high_performance.isChecked() else MAX_WORKERS

            self.worker = OptimizedFolderSearchWorker(
                self.path, raw_keyword, skip_ocr, max_workers
            )
            self.worker.progress.connect(self.progress_bar.setValue)
            self.worker.finished.connect(self.display_summary)
            self.worker.file_processed.connect(self.on_file_found)
            self.worker.start()
        else:
            # Single file search (keeping your original logic but optimized)
            self.search_single_file(raw_keyword, skip_ocr)

    def search_single_file(self, keyword, skip_ocr):
        """Optimized single file search"""
        results = []
        try:
            if self.path.lower().endswith(".pdf"):
                results = self.find_keyword_in_pdf_pages(self.path, keyword, skip_ocr)
                for p in results:
                    self.results_list.addItem(f"Strona {p + 1}")
            elif self.path.lower().endswith(".docx"):
                results = self.find_keyword_in_docx_paragraphs(self.path, keyword)
                for i in results:
                    self.results_list.addItem(f"Paragraph {i + 1}")

            if not results:
                QMessageBox.information(self, "Brak wyników", "Brak wyników w wybranym pliku")
        finally:
            self.search_button.setEnabled(True)
            self.stop_button.setEnabled(False)

    def find_keyword_in_pdf_pages(self, path, keyword, skip_ocr=False):
        """Find specific pages containing keyword"""
        found_pages = []
        norm_keyword = normalize_text(keyword)
        
        try:
            doc = fitz.open(path)
            has_text_pages = False

            for page_num in range(len(doc)):
                text = doc[page_num].get_text()
                if text.strip():
                    has_text_pages = True
                    if norm_keyword in normalize_text(text):
                        found_pages.append(page_num)

            doc.close()

            if not has_text_pages and not skip_ocr:
                if not TESSERACT_CMD:
                    print("[OCR SKIPPED] Tesseract not available")
                    return found_pages
                    
                # OCR search for page numbers
                with TemporaryDirectory() as tmpdir:
                    images = convert_from_path(path, dpi=150, output_folder=tmpdir, poppler_path=POPPLER_PATH)
                    for i, img in enumerate(images):
                        try:
                            text = pytesseract.image_to_string(img, lang='pol')
                            if norm_keyword in normalize_text(text):
                                found_pages.append(i)
                        except Exception as e:
                            print(f"[OCR ERROR]: {e}")
                        finally:
                            img.close()
                            
        except Exception as e:
            print(f"Single-file PDF search error: {e}")
            
        return found_pages

    def find_keyword_in_docx_paragraphs(self, path, keyword):
        """Find specific paragraphs containing keyword"""
        found_paragraphs = []
        norm_keyword = normalize_text(keyword)
        
        try:
            doc = docx.Document(path)
            for i, para in enumerate(doc.paragraphs):
                if norm_keyword in normalize_text(para.text):
                    found_paragraphs.append(i)
        except Exception as e:
            print(f"DOCX search error: {e}")
            
        return found_paragraphs

    def on_file_found(self, filepath, found):
        """Real-time update when file is found"""
        if found:
            self.status_label.setText(f"Znaleziono: {os.path.basename(filepath)}")

    def stop_search(self):
        """Stop the current search"""
        if self.worker and self.worker.isRunning():
            self.worker.terminate()
            self.worker.wait()
        self.progress_bar.setVisible(False)
        self.search_button.setEnabled(True)
        self.stop_button.setEnabled(False)
        self.status_label.setText("Wyszukiwanie zatrzymane")

    def display_summary(self, results, total, skipped):
        self.progress_bar.setVisible(False)
        self.search_button.setEnabled(True)
        self.stop_button.setEnabled(False)
        
        found = len(results)
        for path in results:
            self.results_list.addItem(path)

        if found == 0:
            QMessageBox.information(self, "Brak wyników", "Brak wyników w wybranym folderze")
        else:
            QMessageBox.information(
                self,
                "Raport wyszukiwania",
                f"Pliki przeszukane: {total}\n"
                f"Wyniki znalezione: {found}\n"
                f"Pominięte pliki: {skipped}"
            )
        
        self.status_label.setText(f"Zakończono: {found}/{total} plików ze słowem kluczowym")

    def open_item(self, item):
        text = item.text()
        if self.folder_radio.isChecked():
            # Folder mode: text contains full file path
            open_file(text)
        else:
            # Single file mode: text contains "Page X" or "Paragraph X"
            # Open the selected file (self.path)
            if self.path:
                open_file(self.path)

    def closeEvent(self, event):
        """Clean shutdown"""
        if self.worker and self.worker.isRunning():
            self.worker.terminate()
            self.worker.wait()
        event.accept()

if __name__ == "__main__":
    # Set multiprocessing start method for Windows compatibility
    if platform.system() == "Windows":
        mp.set_start_method('spawn', force=True)
    
    app = QApplication(sys.argv)
    window = KeywordSearcher()
    window.resize(700, 600)
    window.show()
    sys.exit(app.exec_())
